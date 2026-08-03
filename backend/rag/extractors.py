"""여러 문서 형식에서 순수 텍스트를 추출한다.

.txt/.md는 그대로 읽고, .pdf/.docx/.hwpx는 각 형식에 맞는 방식으로 텍스트만 뽑아낸다.
추출된 텍스트는 chunker.py가 청킹할 때 원래 텍스트 파일과 동일하게 취급된다.
"""

from __future__ import annotations

import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import pdfplumber
from docx import Document

PLAIN_TEXT_SUFFIXES = {".txt", ".md"}
SUPPORTED_SUFFIXES = PLAIN_TEXT_SUFFIXES | {".pdf", ".docx", ".hwpx"}


def _local_name(tag: str) -> str:
    """XML 태그에서 네임스페이스를 떼고 이름만 남긴다. 예: '{...}p' -> 'p'."""
    return tag.rsplit("}", 1)[-1]


def _extract_pdf(path: Path) -> str:
    with pdfplumber.open(path) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n\n".join(parts)


def _render_hwpx_table(tbl) -> str:
    """표(hp:tbl) 하나를 사람이 읽을 수 있는 텍스트로 만든다.

    행(hp:tr)은 줄바꿈으로, 셀(hp:tc)은 " | "로 구분한다 — docx 추출기와 같은 형식이다.
    구분자가 없으면 "연번품명물품번호1미니버스25101501..."처럼 값이 전부 붙어버려
    표 안의 숫자를 묻는 질문(기준단가·통계목 등)이 검색되지 않는다.

    행 사이는 빈 줄이 아닌 줄바꿈 하나라서, chunker가 표 하나를 한 문단으로 취급한다.
    """
    rows: list[str] = []
    for tr in tbl:
        if _local_name(tr.tag) != "tr":
            continue
        cells: list[str] = []
        for tc in tr:
            if _local_name(tc.tag) != "tc":
                continue
            # 셀 안에도 문단·표가 들어갈 수 있어(중첩 표) 같은 로직을 재귀로 태운다.
            blocks: list[str] = []
            for child in tc:
                _walk_hwpx(child, blocks)
            cells.append(" ".join(b.replace("\n", " ") for b in blocks).strip())
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _walk_hwpx(node, out: list[str]) -> None:
    """OWPML 트리를 문서 순서대로 훑어 텍스트 블록을 out에 쌓는다.

    표는 sec > p > run > tbl 구조라 바깥 문단 안에 통째로 들어 있다. 그래서 모든 p를
    훑으며 하위 hp:t를 전부 이어붙이면 (1) 표가 구분자 없이 뭉개지고 (2) 셀 안의 p에서
    같은 내용을 또 뽑아 중복된다. 표를 만나면 별도 블록으로 처리하고 더 내려가지 않는다.
    """
    name = _local_name(node.tag)

    if name == "tbl":
        table_text = _render_hwpx_table(node)
        if table_text.strip():
            out.append(table_text)
        return

    if name == "p":
        buffer: list[str] = []

        def visit(elem) -> None:
            for child in elem:
                child_name = _local_name(child.tag)
                if child_name == "tbl":
                    # 표 앞까지의 글을 먼저 내보내 원문 순서를 지킨다.
                    if buffer:
                        text = "".join(buffer)
                        buffer.clear()
                        if text.strip():
                            out.append(text)
                    table_text = _render_hwpx_table(child)
                    if table_text.strip():
                        out.append(table_text)
                elif child_name == "t":
                    if child.text:
                        buffer.append(child.text)
                else:
                    visit(child)

        visit(node)
        if buffer:
            text = "".join(buffer)
            if text.strip():
                out.append(text)
        return

    for child in node:
        _walk_hwpx(child, out)


def _extract_hwpx(path: Path) -> str:
    """HWPX는 OOXML처럼 zip 안에 XML(OWPML)로 저장된 문서다.

    Contents/section*.xml을 문서 순서대로 훑어 문단은 문단대로, 표는 표 형태를 살려 모은다.
    네임스페이스 URI가 버전마다 다를 수 있어 로컬 이름(p, t, tbl 등)만으로 판별한다.
    """
    blocks: list[str] = []
    with zipfile.ZipFile(path) as zf:
        section_names = sorted(
            name for name in zf.namelist()
            if name.startswith("Contents/section") and name.endswith(".xml")
        )
        for name in section_names:
            root = ET.fromstring(zf.read(name))
            for child in root:
                _walk_hwpx(child, blocks)
    return "\n\n".join(blocks)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in PLAIN_TEXT_SUFFIXES:
        return path.read_text(encoding="utf-8")
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".hwpx":
        return _extract_hwpx(path)
    raise ValueError(f"지원하지 않는 파일 형식입니다: {suffix}")
