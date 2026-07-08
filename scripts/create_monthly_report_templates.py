from pathlib import Path

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_PDF = ROOT / "[본문] AI업무비서 월간 운영 결과(3월).pdf"
TMP = ROOT / "tmp" / "pdfs" / "monthly-report-assets"
OUT = ROOT / "output" / "templates"
DOCX_PATH = OUT / "ai-assistant-monthly-report-template.docx"

FONT_BODY = "맑은 고딕"
FONT_SERIF = "함초롬바탕"


def set_run_font(run, name=FONT_BODY, size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, attrs in edges.items():
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), str(value))


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph(paragraph, *, before=0, after=0, line=1.45, align=None, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align
    if keep:
        fmt.keep_with_next = True
    return paragraph


def add_text(paragraph, text, *, size=10.5, bold=False, font=FONT_BODY, color=None):
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold, color)
    return run


def add_heading(document, text):
    p = set_paragraph(document.add_paragraph(), before=9, after=4, line=1.0, keep=True)
    add_text(p, f"□ {text}", size=13, bold=True)
    return p


def add_circle(document, text, *, indent=0.4, after=2):
    p = set_paragraph(document.add_paragraph(), after=after, line=1.5)
    p.paragraph_format.left_indent = Cm(indent)
    add_text(p, f"○ {text}", size=10.5)
    return p


def add_dash(document, text, *, indent=0.85, after=1):
    p = set_paragraph(document.add_paragraph(), after=after, line=1.4)
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    add_text(p, f"- {text}", size=10.2)
    return p


def add_insight(document, text):
    p = set_paragraph(document.add_paragraph(), before=4, after=2, line=1.45)
    p.paragraph_format.left_indent = Cm(0.25)
    add_text(p, f"⇒ {text}", size=10.2)
    return p


def remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                bottom={"val": "nil"},
                start={"val": "nil"},
                end={"val": "nil"},
                insideH={"val": "nil"},
                insideV={"val": "nil"},
            )


def format_cell(cell, text, *, bold=False, size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER, fill=None):
    cell.text = ""
    p = set_paragraph(cell.paragraphs[0], line=1.0, align=align)
    add_text(p, text, size=size, bold=bold, font=FONT_SERIF)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=45, bottom=45, start=55, end=55)
    if fill:
        set_cell_shading(cell, fill)


def extract_pdf_images():
    TMP.mkdir(parents=True, exist_ok=True)
    pdf = fitz.open(SOURCE_PDF)
    names = {
        "city-logo.png": (0, 12),
        "mascot.png": (0, 15),
        "daily-usage.png": (1, 32),
        "department-usage.png": (1, 33),
        "agent-usage.png": (2, 40),
    }
    for name, (_, xref) in names.items():
        pix = fitz.Pixmap(pdf, xref)
        if pix.alpha or pix.n > 4:
            pix = fitz.Pixmap(fitz.csRGB, pix)
        pix.save(TMP / name)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(16)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(5)
    section.footer_distance = Mm(7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(10.5)

    section.footer.paragraphs[0].text = ""


def add_header_block(document):
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Cm(3.2), Cm(10.8), Cm(3.2))
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = width
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    remove_table_borders(table)

    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left.add_run().add_picture(str(TMP / "city-logo.png"), width=Cm(1.55))

    center = table.cell(0, 1).paragraphs[0]
    center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center.paragraph_format.space_before = Pt(5)
    add_text(center, "파  주  시", size=21, bold=True, font=FONT_SERIF)

    right = table.cell(0, 2).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run().add_picture(str(TMP / "mascot.png"), width=Cm(1.55))

    slogan = set_paragraph(document.add_paragraph(), before=1, after=5, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(slogan, "개인정보 방심은 NO, 보안은 YES", size=8.6, font=FONT_SERIF)


def add_approval_table(document):
    table = document.add_table(rows=5, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    widths = [Cm(1.55), Cm(3.05), Cm(1.65), Cm(1.85), Cm(1.85), Cm(1.85), Cm(1.85), Cm(1.85)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    left_values = [
        ("등록번호", "{{문서번호}}"),
        ("등록일자", "{{등록일}}"),
        ("결재일자", "{{결재일}}"),
        ("공개구분", "{{공개구분}}"),
        ("전화번호", "{{전화번호}}"),
    ]
    for idx, (label, value) in enumerate(left_values):
        format_cell(table.cell(idx, 0), label, size=8.2, fill="F5F5F5")
        format_cell(table.cell(idx, 1), value, size=8.1)

    approvers = [("주무관", "{{담당자}}"), ("AI기반팀장", "{{팀장}}"), ("정보통신과장", "{{과장}}")]
    for offset, (role, person) in enumerate(approvers, start=2):
        format_cell(table.cell(0, offset), role, size=7.8, fill="F5F5F5")
        format_cell(table.cell(1, offset), person, size=9.0, bold=True)

    merged = table.cell(0, 5).merge(table.cell(1, 7))
    format_cell(merged, "{{추가결재}}", size=8.2)

    coop_label = table.cell(2, 2).merge(table.cell(4, 2))
    format_cell(coop_label, "협조자", size=8.2, fill="F5F5F5")
    coop_value = table.cell(2, 3).merge(table.cell(4, 7))
    format_cell(coop_value, "{{협조자}}", size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT)

    p = set_paragraph(document.add_paragraph(), after=0, line=1.0)
    p.paragraph_format.space_after = Pt(0)


def add_title_and_summary(document):
    title = set_paragraph(document.add_paragraph(), before=7, after=6, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER, keep=True)
    add_text(title, "AI업무비서 월간 운영 결과({{보고월}})", size=20, bold=True)

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.cell(0, 0).width = Cm(17.2)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF3FB")
    set_cell_margins(cell, top=95, bottom=95, start=120, end=120)
    p = set_paragraph(cell.paragraphs[0], after=1, line=1.38)
    add_text(p, "{{핵심요약}}", size=9.7)
    for token in ("{{요약시사점1}}", "{{요약시사점2}}", "{{요약시사점3}}"):
        p = set_paragraph(cell.add_paragraph(), after=1, line=1.3)
        add_text(p, f"⇒ {token}", size=9.5)


def add_page_one(document):
    add_header_block(document)
    add_approval_table(document)
    add_title_and_summary(document)

    add_heading(document, "목적")
    add_circle(document, "{{운영목적}}")

    p = add_heading(document, "운영기간")
    add_text(p, ":  {{운영기간}}", size=11.5, bold=False)

    add_heading(document, "월간 운영 데이터 분석")
    add_circle(document, "이용 추이 분석", after=1)
    for token in ("{{이용추이1}}", "{{이용추이2}}", "{{이용추이3}}"):
        add_dash(document, token)
    add_insight(document, "{{이용추이시사점}}")


def add_picture_center(document, path, width_cm, *, before=2, after=4):
    p = set_paragraph(document.add_paragraph(), before=before, after=after, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    return p


def add_page_two(document):
    p = set_paragraph(document.add_paragraph(), before=0, after=3, line=1.0)
    add_text(p, "- 이용 그래프 현황", size=10.2)
    add_picture_center(document, TMP / "daily-usage.png", 16.8, after=7)

    add_heading(document, "부서별 이용 현황")
    add_circle(document, "{{부서현황1}}")
    add_circle(document, "{{부서현황2}}")
    add_insight(document, "{{부서시사점}}")
    add_picture_center(document, TMP / "department-usage.png", 17.0, before=5, after=6)

    add_heading(document, "에이전트별 이용 현황")
    add_circle(document, "{{에이전트현황1}}")
    add_circle(document, "{{에이전트현황2}}")
    for token in ("{{에이전트시사점1}}", "{{에이전트시사점2}}", "{{에이전트시사점3}}"):
        add_insight(document, token)


def add_page_three(document):
    add_picture_center(document, TMP / "agent-usage.png", 17.0, before=0, after=7)

    add_heading(document, "특이사항")
    add_circle(document, "{{특이사항1}}")
    for token in ("{{특이원인1}}", "{{특이조치1}}", "{{특이결과1}}"):
        add_dash(document, token)
    add_circle(document, "{{특이사항2}}")
    for token in ("{{특이내용2-1}}", "{{특이내용2-2}}"):
        add_dash(document, token)
    add_circle(document, "{{안정성요약}}")

    add_heading(document, "향후 계획")
    add_circle(document, "{{향후계획}}")

    p = set_paragraph(document.add_paragraph(), before=13, after=2, line=1.5)
    add_text(p, "붙임  1. {{붙임1}}", size=10.5)
    p = set_paragraph(document.add_paragraph(), after=0, line=1.5)
    p.paragraph_format.left_indent = Cm(1.0)
    add_text(p, "2. {{붙임2}}  끝.", size=10.5)


def add_placeholder_reference(document):
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    p = set_paragraph(document.add_paragraph(), after=8, line=1.0)
    add_text(p, "템플릿 입력 필드 안내", size=16, bold=True)
    p = set_paragraph(document.add_paragraph(), after=8, line=1.5)
    add_text(p, "이 페이지는 시스템 등록 시 제거합니다. 중괄호 필드는 웹 입력값 또는 집계 데이터로 교체하며, 그래프 3개는 동일 위치의 이미지 슬롯으로 교체합니다.", size=10.5)

    groups = [
        ("문서·결재", "DOCUMENT_NUMBER, REGISTERED_DATE, APPROVAL_DATE, DISCLOSURE_LEVEL, PHONE_NUMBER, DRAFTER, TEAM_LEADER, DEPARTMENT_HEAD, COOPERATORS"),
        ("기본 정보", "REPORT_MONTH, SUMMARY_SENTENCE, PURPOSE, OPERATION_PERIOD"),
        ("운영 분석", "DAILY_TREND_*, DEPARTMENT_USAGE_*, AGENT_USAGE_*"),
        ("특이사항", "INCIDENT_*, STABILITY_SUMMARY, NEXT_PLAN"),
        ("첨부", "ATTACHMENT_1, ATTACHMENT_2"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_cell(table.cell(0, 0), "구분", bold=True, fill="D9EAF7")
    format_cell(table.cell(0, 1), "필드", bold=True, fill="D9EAF7")
    set_repeat_table_header(table.rows[0])
    for label, fields in groups:
        cells = table.add_row().cells
        format_cell(cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        format_cell(cells[1], fields, size=8.2, align=WD_ALIGN_PARAGRAPH.LEFT)


def create_docx():
    OUT.mkdir(parents=True, exist_ok=True)
    extract_pdf_images()
    document = Document()
    configure_document(document)
    add_page_one(document)
    document.add_page_break()
    add_page_two(document)
    document.add_page_break()
    add_page_three(document)
    document.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    create_docx()
